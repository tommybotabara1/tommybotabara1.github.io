from __future__ import annotations

from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\TJ\OneDrive\Resume_Botabara_Tommy_5262026.docx")
OUTPUT = Path(r"C:\Users\TJ\OneDrive\Resume_Botabara_Tommy_5262026.docx")
SNAPSHOT = Path(r"C:\Users\TJ\OneDrive\Resume_Botabara_Tommy_redo_snapshot_20260704.docx")

NAVY = RGBColor(31, 78, 121)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(88, 88, 88)
LINE = RGBColor(196, 205, 214)


PROFILE = (
    "AI & Technology Professional with 6+ years of experience designing, prototyping, "
    "and deploying AI-powered solutions, enterprise automation, and data platforms across "
    "consulting and software engineering. Hands-on builder of LLM applications and multi-agent "
    "workflows on Claude and Codex, with strengths in rapid prototyping, Python automation, "
    "low-code delivery, stakeholder engagement, and translating business workflows into "
    "production-minded AI automation."
)

EXPERIENCE = [
    {
        "date": "2025 - Present",
        "org": "ARCH GLOBAL SERVICES PHILIPPINES",
        "title": "Software Engineer Level 2",
        "bullets": [
            "Develop and enhance Python ETL pipelines and automation scripts powering enterprise data processes.",
            "Build Snowflake-integrated data workflows with transformations, validation logic, and structured error handling.",
            "Drive modernization initiatives by optimizing SQL queries, refactoring legacy scripts, and applying maintainability and performance best practices.",
            "Serve as point person for Python development, debugging, multi-environment troubleshooting, documentation, and deployment readiness.",
        ],
    },
    {
        "date": "2020 - 2025",
        "org": "DELOITTE TECHNOLOGY & TRANSFORMATION SEA",
        "title": "Senior Consultant - Enterprise Technology and Performance",
        "engagements": [
            (
                "Digital Field Operations - Leading Philippine Telecommunications Provider",
                [
                    "Led a workstream enhancing big-data reconciliation scripts from requirements gathering through deployment and knowledge transfer.",
                    "Built Python scripts implementing 150+ business rules to reconcile customer and inventory data across multiple enterprise databases.",
                    "Assessed data completeness, consistency, and traceability through custom Python validation tooling.",
                ],
            ),
            (
                "Energy Cost Reduction Analysis - Leading Thailand Telco",
                [
                    "Led analytics across 14,000+ cell sites to identify power-consumption outliers and performance-improvement opportunities.",
                    "Built Python scripts that surfaced insights informing the client's energy cost-reduction strategy.",
                ],
            ),
            (
                "Intelligent Business Process Management - Leading Global Food Manufacturer",
                [
                    "Spearheaded a workstream building and enhancing screens in the PIO Claims Management PowerApps application.",
                    "Owned the full lifecycle from learning PowerApps from scratch to production deployment for a business-user-facing prototype.",
                ],
            ),
            (
                "Low-Code Risk Assessment Tool - High-Risk Country Strategy Council",
                [
                    "Led development of a Risk Assessment Tool on Microsoft Power Platform, including Power BI, Power Apps, and Power Automate.",
                    "Designed multi-screen workflows enabling stakeholders to assess country risks affecting global operations.",
                ],
            ),
            (
                "Reporting, Migration, and Vendor Evaluation Workstreams",
                [
                    "Automated Excel and PowerPoint reports using Python, reducing report preparation from days to under a minute.",
                    "Directed data migration for Deloitte SEA's Audit and Assurance practice across 10 countries and moved the team from manual patching to automated scripting.",
                    "Supported enterprise architecture, RFP evaluation, vendor selection, capability mapping, and technical requirements elicitation across automotive, banking, utility, and telco contexts.",
                ],
            ),
        ],
    },
    {
        "date": "2019",
        "org": "INFOR PSSC INC.",
        "title": "Global Financial Controller Java Upgrade - QA & ISO Fixes (Internship)",
        "bullets": [
            "Analyzed security vulnerabilities, delivered code fixes, and performed unit testing.",
        ],
    },
]

AI_PROJECTS = [
    {
        "date": "2026",
        "name": "AIOS Workspace Toolkit - Multi-Domain LLM Agent Operating System",
        "stack": "Claude, Codex, Python, browser automation, Google Workspace",
        "bullets": [
            "Built a reusable AIOS workspace across procurement, real estate, teaching, and personal operations with 25+ workflow agents.",
            "Standardized SKILL.md files, scripts, references, templates, and decision logs for repeatable business-user handoff.",
        ],
    },
    {
        "date": "2026",
        "name": "Airline Procurement Workflow Agents - Supplier and Contract Automation",
        "stack": "Claude, Codex, Microsoft Forms, SAP Ariba, Excel COM, Adobe Sign",
        "bullets": [
            "Prototyped supplier accreditation, contract routing, travel authorization, and structured file-handoff agents across procurement tools.",
        ],
    },
    {
        "date": "2026",
        "name": "Real Estate Brand Growth Site - AI-Assisted Listings and Lead Workflow",
        "stack": "Astro, TypeScript, Google Sheets, Google Drive, automation",
        "bullets": [
            "Created listing, lead-capture, photo prebuild, social-posting, search, upload, and content workflow skills.",
        ],
    },
    {
        "date": "2026",
        "name": "Course Content Pipeline - AI-Assisted Teaching Material Generator",
        "stack": "Markdown, DOCX, Python, Node.js, research workflows",
        "bullets": [
            "Generated source-controlled research packs, module extracts, syllabus updates, assessments, Kahoot imports, and DOCX outputs.",
        ],
    },
    {
        "date": "2026",
        "name": "Wedding Planning Operating System - AI-Assisted Event Automation",
        "stack": "Python, Google Sheets, Apps Script, HTML, planning workflows",
        "bullets": [
            "Connected guest data, seating logic, RSVP setup, floor-plan review, vendor references, and generated event pages.",
        ],
    },
]

SKILLS = [
    ("AI & LLM", "Claude API, OpenAI Codex, LLM agent orchestration, prompt engineering, AI-powered workflow automation, browser automation for AI agents"),
    ("Programming & Data", "Python, Pandas, automation scripting, SQL, Snowflake, MySQL, Jupyter Notebook"),
    ("Low-Code & Enterprise", "Power BI, Power Apps, Power Automate, Adalo, OutSystems, Microsoft 365, Excel COM, Google Workspace"),
    ("Web Development", "HTML, CSS, JavaScript, Django, PHP"),
    ("Tools & Methods", "Visual Studio Code, CLI, Agile and iterative delivery, demo-driven prototyping, stakeholder interviews, technical documentation"),
]

EDUCATION = [
    (
        "2015 - 2019",
        "DE LA SALLE UNIVERSITY",
        "Bachelor of Science in Information Technology",
        "Honorable Mention; Dean's List: AY 2015-16 (T1, T2, T3); AY 2016-17 (T1, T2); AY 2017-18 (T2, T3)",
    ),
    (
        "2011 - 2015",
        "PASAY CITY ACADEMY",
        "High School Diploma - Graduating Class President",
        "Most Outstanding Achiever (2012-2015); Mercury Drug Best in Science and Math (2015)",
    ),
]


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top=40, start=80, bottom=40, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def paragraph_border_bottom(paragraph, color="C4CDD6", size="6", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.05

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet.font.size = Pt(9.15)
    bullet.paragraph_format.left_indent = Inches(1.48)
    bullet.paragraph_format.first_line_indent = Inches(-0.18)
    bullet.paragraph_format.space_after = Pt(1.5)
    bullet.paragraph_format.line_spacing = 1.05


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    set_run_font(run, size=10, bold=True, color=NAVY)
    paragraph_border_bottom(p)
    return p


def add_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TOMMY BOTABARA")
    set_run_font(run, size=20, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("AI & Technology Professional | Software Engineer | Technology Consultant")
    set_run_font(run, size=9.8, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    for text, is_link in [
        ("Mobile: (+63) 917-802-9516 | ", False),
        ("tommybotabara@gmail.com", False),
        (" | LinkedIn: ", False),
        ("linkedin.com/in/tommy-botabara-b13496193", True),
    ]:
        run = p.add_run(text)
        set_run_font(run, size=8.8, color=NAVY if is_link else MUTED)
        if is_link:
            run.underline = True


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        set_run_font(run, size=9.3, color=INK)
    return p


def add_role_header(doc, date, title):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(1.35))
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(date)
    set_run_font(r, size=9.4, bold=True, color=INK)
    p.add_run("\t")
    r = p.add_run(title)
    set_run_font(r, size=9.6, bold=True, color=INK)
    return p


def add_role_subtitle(doc, text, left_indent=1.35):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=9.1, italic=True, color=MUTED)
    return p


def add_bullet(doc, text, left_indent=1.48):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    for run in p.runs:
        set_run_font(run, size=9.0, color=INK)
    return p


def add_engagement(doc, title, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.35)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, size=9.15, bold=True, color=INK)
    for bullet in bullets:
        add_bullet(doc, bullet)


def add_skills_table(doc):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.allow_autofit = False
    widths = [Inches(1.35), Inches(6.0)]
    for label, detail in SKILLS:
        cells = table.add_row().cells
        for idx, width in enumerate(widths):
            cells[idx].width = width
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[idx])
        label_p = cells[0].paragraphs[0]
        label_p.paragraph_format.space_after = Pt(1)
        label_run = label_p.add_run(label)
        set_run_font(label_run, size=9.0, bold=True, color=INK)
        detail_p = cells[1].paragraphs[0]
        detail_p.paragraph_format.space_after = Pt(1)
        detail_run = detail_p.add_run(detail)
        set_run_font(detail_run, size=9.0, color=INK)
    return table


def build_resume():
    if SOURCE.exists() and not SNAPSHOT.exists():
        copyfile(SOURCE, SNAPSHOT)

    doc = Document()
    configure_document(doc)
    add_header(doc)

    add_section_heading(doc, "Profile Summary")
    add_body_paragraph(doc, PROFILE)

    add_section_heading(doc, "Work Experience")
    for item in EXPERIENCE:
        add_role_header(doc, item["date"], item["org"])
        add_role_subtitle(doc, item["title"])
        if "bullets" in item:
            for bullet in item["bullets"]:
                add_bullet(doc, bullet)
        for title, bullets in item.get("engagements", []):
            add_engagement(doc, title, bullets)

    add_section_heading(doc, "AI Projects")
    for project in AI_PROJECTS:
        add_role_header(doc, project["date"], project["name"])
        add_role_subtitle(doc, f"Self-directed; {project['stack']}")
        for bullet in project["bullets"]:
            add_bullet(doc, bullet)

    add_section_heading(doc, "Technical Skills")
    add_skills_table(doc)

    add_section_heading(doc, "Education")
    for date, school, degree, honors in EDUCATION:
        add_role_header(doc, date, school)
        add_role_subtitle(doc, degree)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(1.35)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(honors)
        set_run_font(r, size=9.0, italic=True, color=MUTED)

    doc.save(OUTPUT)
    print(OUTPUT)
    print(SNAPSHOT)


if __name__ == "__main__":
    build_resume()
